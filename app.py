
from fastapi import FastAPI, HTTPException, Request, UploadFile, File, Form
from fastapi.responses import FileResponse, HTMLResponse, RedirectResponse
from pydantic import BaseModel, Field
from typing import List, Optional, Tuple
from pathlib import Path
from copy import deepcopy
from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from PIL import Image, ImageOps
import base64, uuid, re, requests, urllib.parse

BASE_DIR = Path(__file__).resolve().parent
TEMPLATE_PATH = BASE_DIR / "template.docx"
OUT_DIR = BASE_DIR / "outputs"
UPLOAD_DIR = BASE_DIR / "uploads"
OUT_DIR.mkdir(exist_ok=True)
UPLOAD_DIR.mkdir(exist_ok=True)

app = FastAPI(title="Receipt DOCX Generator", version="1.0.10")

class OpenAIFileRef(BaseModel):
    name: Optional[str] = None
    id: Optional[str] = None
    mime_type: Optional[str] = None
    download_link: Optional[str] = None

class ReceiptItem(BaseModel):
    use_date: str
    store: str
    amount: str
    receipt_image_url: Optional[str] = None
    receipt_image_base64: Optional[str] = None
    receipt_image_file_path: Optional[str] = None

class ReceiptRequest(BaseModel):
    receipts: List[ReceiptItem]
    openaiFileIdRefs: Optional[List[OpenAIFileRef]] = None
    card_name: str = "중앙청년지원센터 법인카드(0000)"
    user_name: str = "중앙청년지원센터 매니저"
    purpose: str = ""
    output_filename: Optional[str] = "법인카드_영수증_증빙자료.docx"

def clear_cell(cell):
    tc = cell._tc
    for child in list(tc):
        if child.tag != qn("w:tcPr"):
            tc.remove(child)

def style_run(run, size_pt=13, bold=False):
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(size_pt)
    run.font.bold = bold

def write_center(cell, text, size_pt=13, bold=False):
    clear_cell(cell)
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text or "")
    style_run(run, size_pt=size_pt, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def write_amount_cell(cell, amount):
    clear_cell(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p1 = cell.add_paragraph(); p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(amount or ""); style_run(r1, 13)
    p2 = cell.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("* 비목: 인증연구 > 업무추진비 > 회의비"); style_run(r2, 11)

def image_from_base64(s):
    data = s.strip()
    if "," in data and data.lower().startswith("data:image"):
        data = data.split(",", 1)[1]
    data = re.sub(r"\s+", "", data)
    if len(data) % 4:
        data += "=" * (4 - len(data) % 4)
    img = Image.open(BytesIO(base64.b64decode(data, validate=False)))
    img = ImageOps.exif_transpose(img); img.load()
    return img

def image_from_url(url):
    resp = requests.get(url, timeout=30)
    resp.raise_for_status()
    img = Image.open(BytesIO(resp.content))
    img = ImageOps.exif_transpose(img); img.load()
    return img

def save_img(img, path):
    if img.mode not in ("RGB", "RGBA"):
        img = img.convert("RGB")
    img.save(path, format="PNG")
    return path

def download_link_from_refs(refs, index=0):
    if not refs:
        return None
    image_refs = []
    for r in refs:
        mime = (r.mime_type or "").lower()
        name = (r.name or "").lower()
        if r.download_link and (mime.startswith("image/") or name.endswith((".png",".jpg",".jpeg",".webp",".bmp"))):
            image_refs.append(r)
    if image_refs:
        return image_refs[min(index, len(image_refs)-1)].download_link
    links = [r for r in refs if r.download_link]
    return links[min(index, len(links)-1)].download_link if links else None

def save_receipt_image(receipt, output_path, refs, index):
    img, status = None, "no_image_source"
    dl = download_link_from_refs(refs, index)
    if dl:
        try:
            img = image_from_url(dl); status = "image_loaded_from_openaiFileIdRefs_download_link"
        except Exception as e:
            status = f"openaiFileIdRefs_download_link_error:{type(e).__name__}"
    if img is None and receipt.receipt_image_url:
        try:
            img = image_from_url(receipt.receipt_image_url); status = "image_loaded_from_url"
        except Exception as e:
            status = f"url_error:{type(e).__name__}"
    if img is None and receipt.receipt_image_file_path:
        val = receipt.receipt_image_file_path
        if val.startswith("http://") or val.startswith("https://"):
            try:
                img = image_from_url(val); status = "image_loaded_from_file_path_url"
            except Exception as e:
                status = f"file_path_url_error:{type(e).__name__}"
        else:
            status = f"file_path_not_accessible:{val}"
    if img is None and receipt.receipt_image_base64:
        try:
            img = image_from_base64(receipt.receipt_image_base64); status = "image_loaded_from_base64"
        except Exception as e:
            status = f"base64_error:{type(e).__name__}"
    if img is None:
        return None, status
    return save_img(img, output_path), status

def insert_receipt_area(doc, img_path):
    table = doc.tables[0]
    cell = table.cell(7, 0)
    clear_cell(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if img_path is None:
        return
    with Image.open(img_path) as im:
        w_px, h_px = im.size
    max_w, max_h = 5.45, 3.25
    aspect = w_px / h_px if h_px else 1
    width = min(max_w, max_h * aspect)
    height = width / aspect
    if height > max_h:
        height = max_h
        width = height * aspect
    p.add_run().add_picture(str(img_path), width=Inches(width))

def fill_template(receipt, card_name, user_name, purpose, img_path):
    doc = Document(str(TEMPLATE_PATH))
    table = doc.tables[0]
    write_center(table.cell(1,1), card_name, 13)
    write_center(table.cell(1,3), receipt.use_date, 13)
    write_center(table.cell(2,1), user_name, 13)
    write_center(table.cell(3,1), purpose or "", 13)
    write_center(table.cell(4,1), receipt.store, 13)
    write_amount_cell(table.cell(5,1), receipt.amount)
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    insert_receipt_area(doc, img_path)
    return doc

def append_doc_body(target, source):
    target.add_page_break()
    body = target.element.body
    for child in source.element.body:
        if not child.tag.endswith("sectPr"):
            body.append(deepcopy(child))

def save_final_doc(request, pages, filename):
    folder_id = str(uuid.uuid4())
    folder = OUT_DIR / folder_id
    folder.mkdir(parents=True, exist_ok=True)
    final_doc = pages[0]
    for page in pages[1:]:
        append_doc_body(final_doc, page)
    if not filename.endswith(".docx"):
        filename += ".docx"
    out_path = folder / filename
    final_doc.save(out_path)
    encoded = urllib.parse.quote(filename)
    base_url = str(request.base_url).rstrip("/")
    return f"{base_url}/download/{folder_id}/{encoded}", filename

@app.get("/", response_class=HTMLResponse)
def root():
    return "<h2>Receipt DOCX Generator v1.0.10</h2><p><a href='/form'>form</a> | <a href='/bulk-form'>bulk-form</a> | <a href='/health'>health</a></p>"

@app.get("/health")
def health():
    return {"status":"ok","version":"1.0.10"}

@app.get("/form", response_class=HTMLResponse)
def form_page():
    return """
    <html><body style='font-family:sans-serif; padding:40px;'>
    <h2>영수증 증빙자료 DOCX 생성</h2>
    <form action="/create-from-form" enctype="multipart/form-data" method="post">
      <p>사용일자<br><input name="use_date" required></p>
      <p>사용처<br><input name="store" required></p>
      <p>사용금액<br><input name="amount" required></p>
      <p>영수증 이미지<br><input name="receipt_image" type="file" accept="image/*" required></p>
      <button type="submit">DOCX 생성</button>
    </form></body></html>
    """

@app.post("/create-from-form")
async def create_from_form(request: Request, use_date: str = Form(...), store: str = Form(...), amount: str = Form(...), receipt_image: UploadFile = File(...)):
    folder = UPLOAD_DIR / str(uuid.uuid4())
    folder.mkdir(parents=True, exist_ok=True)
    img_path = folder / "receipt.png"
    img = Image.open(receipt_image.file)
    img = ImageOps.exif_transpose(img); img.load()
    save_img(img, img_path)
    receipt = ReceiptItem(use_date=use_date, store=store, amount=amount)
    page = fill_template(receipt, "중앙청년지원센터 법인카드(0000)", "중앙청년지원센터 매니저", "", img_path)
    url, _ = save_final_doc(request, [page], "법인카드_영수증_증빙자료.docx")
    return RedirectResponse(url, status_code=303)

@app.get("/bulk-form", response_class=HTMLResponse)
def bulk_form_page():
    fields = ""
    for i in range(1, 11):
        req = "required" if i == 1 else ""
        fields += f"<fieldset><legend>영수증 {i}</legend><p>사용일자<br><input name='use_date_{i}' {req}></p><p>사용처<br><input name='store_{i}' {req}></p><p>사용금액<br><input name='amount_{i}' {req}></p><p>영수증 이미지<br><input name='receipt_image_{i}' type='file' accept='image/*' {req}></p></fieldset>"
    return f"<html><body style='font-family:sans-serif; padding:40px;'><h2>여러 영수증 DOCX 생성</h2><form action='/create-bulk-from-form' enctype='multipart/form-data' method='post'>{fields}<button type='submit'>통합 DOCX 생성</button></form></body></html>"

@app.post("/create-bulk-from-form")
async def create_bulk_from_form(request: Request):
    form = await request.form()
    pages = []
    folder = UPLOAD_DIR / str(uuid.uuid4()); folder.mkdir(parents=True, exist_ok=True)
    for i in range(1, 11):
        use_date = (form.get(f"use_date_{i}") or "").strip()
        store = (form.get(f"store_{i}") or "").strip()
        amount = (form.get(f"amount_{i}") or "").strip()
        image = form.get(f"receipt_image_{i}")
        if not use_date and not store and not amount:
            continue
        if not use_date or not store or not amount or not image:
            raise HTTPException(status_code=400, detail=f"영수증 {i} 필수 정보 누락")
        img_path = folder / f"receipt_{i}.png"
        img = Image.open(image.file); img = ImageOps.exif_transpose(img); img.load()
        save_img(img, img_path)
        receipt = ReceiptItem(use_date=use_date, store=store, amount=amount)
        pages.append(fill_template(receipt, "중앙청년지원센터 법인카드(0000)", "중앙청년지원센터 매니저", "", img_path))
    if not pages:
        raise HTTPException(status_code=400, detail="입력된 영수증이 없습니다.")
    url, _ = save_final_doc(request, pages, "법인카드_영수증_증빙자료.docx")
    return RedirectResponse(url, status_code=303)

@app.get("/download/{folder_id}/{filename}")
def download_file(folder_id: str, filename: str):
    file_path = OUT_DIR / Path(folder_id).name / Path(filename).name
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="파일을 찾을 수 없습니다.")
    return FileResponse(file_path, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename=Path(filename).name)

@app.post("/create-receipt-docx")
def create_receipt_docx(req: ReceiptRequest, request: Request):
    if not TEMPLATE_PATH.exists():
        raise HTTPException(status_code=500, detail="template.docx 파일이 서버에 없습니다.")
    if not req.receipts:
        raise HTTPException(status_code=400, detail="영수증 데이터가 없습니다.")
    tmp = OUT_DIR / str(uuid.uuid4()); tmp.mkdir(parents=True, exist_ok=True)
    pages, statuses = [], []
    for i, receipt in enumerate(req.receipts):
        img_path = tmp / f"receipt_{i+1}.png"
        saved, status = save_receipt_image(receipt, img_path, req.openaiFileIdRefs, i)
        statuses.append(status)
        pages.append(fill_template(receipt, req.card_name, req.user_name, req.purpose, saved))
    url, filename = save_final_doc(request, pages, req.output_filename or "법인카드_영수증_증빙자료.docx")
    return {"status":"success", "message":"DOCX 파일이 생성되었습니다.", "download_url":url, "filename":filename, "image_status":" | ".join(statuses)}
